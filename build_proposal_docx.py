#!/usr/bin/env python3
"""
Build the Veldora proposal .docx from its markdown source.

Styling is reproduced from Rev A (measured out of the original file, which was
built before this script existed) so Rev B is visually identical:

  page      Letter, 0.87" margins
  Normal    Calibri 10.5
  H1        Calibri 18   bold  #1F4E79      (used for "Contents" only)
  H2        Calibri 13.5 bold  #1F4E79      (numbered sections)
  H3        Calibri 11.5 bold  #595959
  tables    9pt, header row bold, alternating #D3DFEE banding, first column
            bold, #4F81BD borders, FIXED layout with an explicit tblGrid
  contents  right tab stop at 9411 twips with a dot leader

Two things are load-bearing and easy to lose:
  * `tblLayout fixed` AND an explicit `tblGrid` are both required, or Word
    re-flows the column widths and the wide tables become unreadable.
  * The contents-table regex must not treat the markdown separator row as a
    row of data — `| | |` matches a separator pattern too.

    python3 build_proposal_docx.py <input.md> <output.docx>
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)
BAND = 'D3DFEE'
BORDER = '4F81BD'
BODY_FONT = 'Calibri'
MONO_FONT = 'Consolas'
CONTENT_WIDTH_TWIPS = 10530  # 8.5in - 2 x 0.87in margins, in twentieths of a point


# ── inline markdown ─────────────────────────────────────────────────────────
def add_runs(paragraph, text, size=None, italic=False):
    """Render **bold** and `code` spans into a paragraph."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            # Consolas runs visually larger than Calibri at the same point size,
            # so inline code is stepped down to sit level with the prose.
            run.font.size = Pt((size or 10.5) - 1)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = BODY_FONT if run.font.name != MONO_FONT else MONO_FONT
        if size:
            run.font.size = Pt(size)
        run.italic = italic or None


# ── tables ──────────────────────────────────────────────────────────────────
def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_separator(cells):
    return all(re.fullmatch(r':?-+:?', c) for c in cells if c)


def set_cell_fill(cell, colour):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), colour)
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_borders(table):
    borders = OxmlElement('w:tblBorders')
    for edge, size in (('top', 8), ('left', 8), ('bottom', 8), ('right', 8),
                       ('insideH', 4), ('insideV', 4)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), BORDER)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def set_fixed_layout(table, widths):
    """Fixed layout + an explicit grid. Both, or Word re-flows the columns."""
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(layout)

    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w)))
        grid.append(col)
    table._tbl.insert(list(table._tbl).index(table._tbl.tblPr) + 1, grid)

    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Pt(int(w) / 20)


def no_row_splitting(table):
    """
    Keep each row whole on one page.

    Without this a tall row straddling a page break is torn in half: the wide
    cell's overflow continues on the next page while the narrow label cells beside
    it render empty, which reads as a missing entry rather than a continuation.
    Word then moves the whole row down instead.
    """
    for row in table.rows:
        el = OxmlElement('w:cantSplit')
        row._tr.get_or_add_trPr().append(el)


def repeat_header(table):
    trPr = table.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def column_widths(rows):
    """
    Proportional to content, but never narrower than the longest unbreakable word.

    Without the floor a column sharing a table with a long prose column is
    squeezed until its own header wraps one letter per line — "Week" rendered as
    W/e/e/k down the page, which is what the first build did.
    """
    n = len(rows[0])
    def longest_word(i):
        return max((len(w) for r in rows for w in re.split(r'\s+', r[i]) if w), default=1)
    def longest_cell(i):
        return max(len(r[i]) for r in rows)

    # ~110 twips per character at 9pt Calibri, plus cell padding.
    mins = [max(620, longest_word(i) * 110 + 220) for i in range(n)]
    if sum(mins) >= CONTENT_WIDTH_TWIPS:
        scale = CONTENT_WIDTH_TWIPS / sum(mins)
        return [m * scale for m in mins]

    weights = [max(longest_cell(i), 4) for i in range(n)]
    free = CONTENT_WIDTH_TWIPS - sum(mins)
    total = sum(weights)
    return [mins[i] + free * weights[i] / total for i in range(n)]


def add_table(doc, rows, header, size=9.0, bold_first_col=True, aligns=None):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    body = rows[1:] if header else rows
    for i, source in enumerate(rows):
        cells = table.add_row().cells
        is_header = header and i == 0
        # Banding starts on the first BODY row, so a headerless table (the cover
        # block) bands from its own first row.
        band = ((i - 1) % 2 == 0) if header else (i % 2 == 0)
        for j, text in enumerate(source):
            cell = cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_runs(para, text, size=size)
            if aligns and aligns[j] == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if is_header or (bold_first_col and j == 0):
                    run.bold = True
            if band and not is_header:
                set_cell_fill(cell, BAND)

    set_fixed_layout(table, column_widths(rows))
    no_row_splitting(table)
    if header:
        repeat_header(table)
    doc.add_paragraph()
    return table


# ── document skeleton ───────────────────────────────────────────────────────
def configure_styles(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
        setattr(section, attr, Inches(0.87))

    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, colour in (('Heading 1', 18, NAVY),
                               ('Heading 2', 13.5, NAVY),
                               ('Heading 3', 11.5, GREY)):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(14 if name != 'Heading 3' else 10)
        st.paragraph_format.space_after = Pt(6)


def centered(doc, text, size, colour, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.bold = bold or None
    run.italic = italic or None
    return p


def contents_line(doc, number, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # 9411 twips, matching Rev A — short of the right margin so the page number
    # sits clear of the edge.
    p.paragraph_format.tab_stops.add_tab_stop(
        Pt(9411 / 20), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(f'{number}. {title}\t{page}')
    run.font.name = BODY_FONT
    run.font.size = Pt(10.5)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ── the build ───────────────────────────────────────────────────────────────
def build(md_path, out_path):
    lines = open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    configure_styles(doc)

    i = 0
    state = 'cover'
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- fenced block (the pipeline diagram) --------------------------
        if stripped.startswith('```'):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run('\n'.join(block))
            run.font.name = MONO_FONT
            run.font.size = Pt(8.5)
            continue

        # ---- tables ------------------------------------------------------
        if stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(split_row(lines[i]))
                i += 1
            # A separator is only a separator when it is the SECOND row; the
            # cover block's own first row is `| | |`, which matches otherwise.
            has_header = len(rows) > 1 and is_separator(rows[1]) and any(rows[0])
            aligns = None
            if len(rows) > 1 and is_separator(rows[1]):
                # `:-:` asks for a centred column — the tick/dash matrix in §3
                # is unreadable left-aligned.
                aligns = ['center' if c.startswith(':') and c.endswith(':') else 'left'
                          for c in rows[1]]
                del rows[1]
            # `| | |` above a separator is markdown's headerless-table idiom. Kept,
            # it renders as a banded empty strip across the top of the cover block.
            if rows and not any(c.strip() for c in rows[0]):
                del rows[0]
            if state == 'contents':
                for row in rows:
                    if len(row) == 3 and row[0].isdigit():
                        contents_line(doc, row[0], row[1], row[2])
                continue
            add_table(doc, rows, header=has_header,
                      size=10.0 if state == 'cover' else 9.0, aligns=aligns)
            continue

        # ---- headings ----------------------------------------------------
        if stripped.startswith('### '):
            doc.add_paragraph(stripped[4:], style='Heading 3')
            i += 1
            continue
        if stripped.startswith('## '):
            title = stripped[3:]
            if state == 'cover' and title.startswith('Weather'):
                centered(doc, title, 14, GREY)
                i += 1
                continue
            if title == 'Contents':
                page_break(doc)
                doc.add_paragraph(title, style='Heading 1')
                state = 'contents'
                i += 1
                continue
            if state == 'contents':
                page_break(doc)
                state = 'body'
            doc.add_paragraph(title, style='Heading 2')
            i += 1
            continue
        if stripped.startswith('# '):
            centered(doc, 'VELDORA STUDIO', 13, GREY, bold=True)
            centered(doc, stripped[2:], 26, NAVY, bold=True)
            i += 1
            continue

        # ---- rules, bullets, numbers, prose -------------------------------
        if stripped in ('---', ''):
            i += 1
            continue
        if stripped.startswith('- '):
            text = stripped[2:]
            while i + 1 < len(lines) and lines[i + 1].startswith('  ') and lines[i + 1].strip():
                i += 1
                text += ' ' + lines[i].strip()
            add_runs(doc.add_paragraph(style='List Bullet'), text)
            i += 1
            continue
        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            while i + 1 < len(lines) and lines[i + 1].startswith('   ') and lines[i + 1].strip():
                i += 1
                text += ' ' + lines[i].strip()
            add_runs(doc.add_paragraph(style='List Number'), text)
            i += 1
            continue

        # a paragraph is every following non-blank, non-structural line
        text = stripped
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if not nxt or nxt.startswith(('#', '|', '- ', '---', '```')) or re.match(r'^\d+\. ', nxt):
                break
            i += 1
            text += ' ' + lines[i].strip()
        if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
            centered(doc, text.strip('*'), 10, GREY, italic=True)
        else:
            add_runs(doc.add_paragraph(), text)
        i += 1

    doc.save(out_path)
    print(f'wrote {out_path}')


if __name__ == '__main__':
    build(sys.argv[1], sys.argv[2])
